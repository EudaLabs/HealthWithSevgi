#!/usr/bin/env python3
"""
Generate the filled ISO 42001 Report (Ch 1-3) from the template.
Uses the original .docx template and fills in project-specific content.
"""
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.oxml.ns import qn

TEMPLATE = Path(__file__).parent / "ISO42001_Report_Template_V4_FINAL (1) (3).docx"
OUTPUT   = Path(__file__).parent / "ISO42001_Report_Ch1-3_FILLED.docx"


# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_text(cell, text, bold=False, italic=False, size=None):
    """Clear a cell and write new text, preserving first-run font style."""
    # Keep existing paragraph (preserves alignment/spacing)
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    # Write into first paragraph
    p = cell.paragraphs[0]
    if p.runs:
        run = p.runs[0]
    else:
        run = p.add_run()
    run.text = text
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)


def set_cell_rich(cell, parts):
    """Write rich text into a cell. parts = [(text, bold, italic), ...]"""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    p = cell.paragraphs[0]
    # Clear existing runs
    for run in p.runs:
        run._element.getparent().remove(run._element)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic


def add_table_row(table, cells_data, bold_first=False):
    """Add a row to table. cells_data = list of strings."""
    row = table.add_row()
    for i, txt in enumerate(cells_data):
        cell = row.cells[i]
        for p in cell.paragraphs:
            p.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(txt))
        if bold_first and i == 0:
            run.bold = True
        run.font.size = Pt(9)
    return row


def copy_row_style(source_row, target_row):
    """Copy shading from source row cells to target row cells."""
    for src_cell, tgt_cell in zip(source_row.cells, target_row.cells):
        src_shd = src_cell._tc.find(qn('w:tcPr'))
        if src_shd is not None:
            src_shd_elem = src_shd.find(qn('w:shd'))
            if src_shd_elem is not None:
                tgt_pr = tgt_cell._tc.find(qn('w:tcPr'))
                if tgt_pr is None:
                    tgt_pr = deepcopy(src_shd)
                    tgt_cell._tc.insert(0, tgt_pr)


# ── main ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document(str(TEMPLATE))
    tables = doc.tables

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # T00: Team Info (3x2)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    t = tables[0]
    set_cell_text(t.rows[0].cells[1], "HealthWithSevgi", italic=True)
    set_cell_text(t.rows[1].cells[1],
        "Batuhan Bayazıt (202228008), Efe Çelik (202128016), "
        "Berat Mert Gökkaya (202228019), Burak Aydoğmuş (202128028), "
        "Berfin Duru Alkan (202228005)", italic=True)
    set_cell_text(t.rows[2].cells[1], "Dr. Sevgi Koyuncu Tunç")

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # T05: 1.1 AI System Description (6x2)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    t = tables[5]
    data_11 = [
        "HealthWithSevgi — ML Visualization Tool for Healthcare Professionals",

        "HealthWithSevgi is a browser-based educational tool that enables healthcare "
        "professionals and medical students to learn machine learning concepts through a "
        "guided 7-step wizard pipeline (Clinical Context, Data Exploration, Data Preparation, "
        "Model Selection & Parameters, Results Evaluation, Explainability via SHAP, and Ethics "
        "& Bias Assessment) using real clinical datasets across 20 medical specialties. It does "
        "not make clinical decisions — it teaches ML literacy in a healthcare context.",

        "8 supervised classification algorithms: K-Nearest Neighbors (KNN), Support Vector "
        "Machine (SVM), Decision Tree, Random Forest, Logistic Regression, Naive Bayes, "
        "XGBoost, and LightGBM. Supporting techniques: SMOTE for class imbalance correction, "
        "SHAP (TreeExplainer, LinearExplainer, KernelExplainer) for model explainability, and "
        "subgroup fairness analysis for bias detection.",

        "Browser-based SPA (React 18 + FastAPI), Docker-containerised for HuggingFace Spaces "
        "deployment. Educational use only — no clinical decision-making. All data processing "
        "occurs in-memory with no persistent database.",

        "Design, development, and educational deployment — NOT clinical production. Course "
        "project lifecycle aligned with SENG 430 (5 sprints, February–May 2026).",
    ]
    for i, txt in enumerate(data_11):
        set_cell_text(t.rows[i + 1].cells[1], txt)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # T06: 1.2 AIMS Scope Statement (2x2)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    t = tables[6]

    in_scope = (
        "• Pipeline Steps: All 7 steps — Clinical Context, Data Exploration, Data Preparation "
        "(train/test split, missing value handling, normalisation, SMOTE), Model Selection & "
        "Hyperparameter Tuning, Results & Metrics Evaluation, Explainability (SHAP), Ethics & "
        "Bias Assessment + EU AI Act Compliance\n"
        "• Clinical Domains: 20 medical specialties — Cardiology (Heart Failure, Stroke, "
        "Arrhythmia), Radiology, Nephrology, Oncology (Breast Cancer, Cervical Cancer), "
        "Neurology, Endocrinology (Diabetes, Thyroid), Hepatology, Mental Health, Pulmonology, "
        "Haematology, Dermatology, Ophthalmology, Orthopaedics, ICU/Sepsis, Obstetrics, "
        "Pharmacy\n"
        "• ML Models: 8 classification algorithms (KNN, SVM, Decision Tree, Random Forest, "
        "Logistic Regression, Naive Bayes, XGBoost, LightGBM) with hyperparameter tuning via "
        "RandomizedSearchCV (20 iterations, 3-fold stratified CV)\n"
        "• Users: Healthcare professionals, medical/nursing students, SENG 430 course "
        "participants and instructors\n"
        "• Lifecycle: Design through educational deployment on HuggingFace Spaces\n"
        "• Data: 20 published, de-identified clinical datasets (UCI, Kaggle, PhysioNet, "
        "Harvard Dataverse) + user-uploaded CSV files\n"
        "• Explainability: SHAP-based global feature importance and single-patient waterfall "
        "with 251 clinical name mappings\n"
        "• Ethics: Subgroup fairness audit (gender, age bands), 10% sensitivity gap threshold, "
        "EU AI Act checklist (8 items), 3 real-world AI failure case studies"
    )

    out_scope = (
        "• Clinical diagnosis or treatment decisions — WHY: Educational tool only; no clinical "
        "validation performed; clinical disclaimers displayed in Steps 6 and 7\n"
        "• Real patient-identifiable data (PII/PHI) — WHY: Privacy by design; only published "
        "de-identified datasets; in-memory processing; no server persistence\n"
        "• Production hospital deployment — WHY: No regulatory approval (CE marking, FDA); no "
        "prospective clinical validation study\n"
        "• Post-deployment production monitoring — WHY: Educational lifecycle ends May 2026; "
        "theoretical monitoring plan in Ch. 5.3 (final submission)\n"
        "• Deep learning / neural networks — WHY: Out of pedagogical scope; tool focuses on "
        "interpretable classical ML for SHAP-based explainability\n"
        "• Real-time streaming data or IoT — WHY: Not applicable to educational CSV-based "
        "analysis workflow\n"
        "• Regression or unsupervised learning — WHY: Scope limited to supervised "
        "classification for focused learning objectives\n"
        "• Multi-language / internationalisation — WHY: English-only; single-locale deployment "
        "for university course context"
    )

    set_cell_text(t.rows[0].cells[1], in_scope)
    set_cell_text(t.rows[1].cells[1], out_scope)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # T07: 1.3 Internal & External Issues (9x2)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    t = tables[7]
    issues = [
        ("External: Regulatory — EU AI Act",
         "EU AI Act (Regulation 2024/1689) classifies healthcare AI as high-risk under "
         "Article 6 and Annex III. Though educational, Step 7 includes an 8-item EU AI Act "
         "compliance checklist covering explainability, data documentation, bias auditing, "
         "human oversight, GDPR, monitoring, incident reporting, and clinical validation."),

        ("External: Regulatory — GDPR/KVKK",
         "Turkish KVKK (Law No. 6698) and EU GDPR require data protection impact assessments "
         "for health data. Mitigated by processing only published, de-identified datasets with "
         "no persistent storage. In-memory session architecture with automatic LRU eviction."),

        ("External: Societal — Public concern about healthcare AI",
         "Growing distrust fuelled by documented failures: racially biased pulse oximeters "
         "(2020), over-alerting sepsis systems (2021), dermatology AI underperforming on dark "
         "skin (2019). Step 7 includes these as educational case studies."),

        ("External: Technological — Rapid evolution of XAI",
         "XAI methods evolve rapidly (SHAP, LIME, counterfactual explanations). Our SHAP "
         "implementation covers TreeExplainer, LinearExplainer, and KernelExplainer. Modular "
         "ExplainService architecture allows future integration of alternatives."),

        ("External: Market — Demand for AI-literate clinicians",
         "WHO Global Strategy on Digital Health and EU policy emphasise AI literacy for "
         "healthcare workers. Our tool serves this need across 20 specialties."),

        ("Internal: Competence — Team ML/AI experience",
         "5 software engineering students with varying ML expertise. Mitigated by: structured "
         "7-step wizard, well-established libraries (scikit-learn, XGBoost, LightGBM, SHAP), "
         "iterative development across 5 sprints with instructor feedback."),

        ("Internal: Infrastructure — No GPU, browser-based",
         "All ML computation on CPU via FastAPI. Docker containerisation (python:3.12-slim) "
         "for HuggingFace Spaces. OrderedDict-based LRU eviction (max 50 sessions/models)."),

        ("Internal: Timeline — 10-week sprint cycle",
         "5 sprints with competing coursework. Mitigated by progressive delivery: Sprint 1 "
         "(Steps 1–2), Sprint 2 (Steps 3–4), Sprint 3 (Steps 5–6), Sprint 4 (Step 7 + "
         "Docker), Sprint 5 (user testing + ISO report + presentation)."),

        ("Internal: Resource — In-memory architecture",
         "No external database simplifies privacy compliance but limits capacity. MLService "
         "uses OrderedDict LRU eviction (max 50 sessions/models). DataService uses unbounded "
         "dict, mitigated by sessions being transferred to MLService's LRU store; orphaned "
         "DataService entries are garbage-collected. Production would require LRU cap on "
         "DataService (identified as risk item for Ch. 4)."),
    ]
    # Fill existing rows, add extra if needed
    for i, (issue, desc) in enumerate(issues):
        if i + 1 < len(t.rows):
            row = t.rows[i + 1]
        else:
            row = add_table_row(t, ["", ""])
            copy_row_style(t.rows[1], row)
        set_cell_text(row.cells[0], issue, bold=True)
        set_cell_text(row.cells[1], desc)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # T08: 1.4 Stakeholder Register (8x5)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    t = tables[8]
    stakeholders = [
        ("Healthcare Professionals", "End User",
         "Clinical language; trust model fairness; data privacy; understand AI limitations",
         "No-code UI; Step 6 SHAP with 251 clinical name mappings; Step 7 subgroup bias "
         "audit; privacy notice; clinical disclaimers",
         "High"),
        ("Patients (indirect)", "Affected Party",
         "Fair treatment; data not misused; AI not replacing clinical judgment",
         "Step 7 subgroup analysis (gender, age bands); 10% sensitivity gap threshold; "
         "SMOTE; educational-only scope; no PII processing",
         "High"),
        ("Students / Dev Team\n(5 members)", "Developer",
         "Clear requirements; learn responsible AI; feasible workload",
         "User Guide as spec; ISO deliverables as learning; 5-sprint timeline; Jira-tracked "
         "task distribution across backend, frontend, and QA",
         "Medium"),
        ("Instructor\n(Dr. Sevgi Koyuncu Tunç)", "Governance",
         "Quality; curriculum alignment; ethical AI education",
         "Sprint gates as management review; ISO rubric; peer audit; PDCA cycle documented",
         "High"),
        ("Regulators\n(EU AI Act)", "External",
         "Transparency; human oversight; bias audit; documentation",
         "Step 7 EU AI Act checklist (8 items); SHAP explainability; clinical disclaimers; "
         "all ISO deliverables as governance evidence",
         "High"),
        ("Dataset Providers\n(UCI, Kaggle, etc.)", "Supplier",
         "Proper attribution; ethical use; license compliance",
         "Provenance documentation in Section 3.1; educational-use scope; license compliance "
         "in Section 3.4",
         "Low"),
        ("General Public", "Societal",
         "Trust in healthcare AI; transparency about capabilities and limitations",
         "Open educational tool; Step 7 AI failure case studies; responsible AI principles "
         "embedded throughout; fairness analysis promotes equity",
         "Low"),
    ]
    # Template has 7 data rows; if we need 8th, add extra row
    if len(stakeholders) >= len(t.rows) - 1:
        pass  # 7 stakeholders fit in 7 data rows
    for i, (name, role, needs, response, influence) in enumerate(stakeholders):
        set_cell_text(t.rows[i + 1].cells[0], name)
        set_cell_text(t.rows[i + 1].cells[1], role)
        set_cell_text(t.rows[i + 1].cells[2], needs)
        set_cell_text(t.rows[i + 1].cells[3], response)
        set_cell_text(t.rows[i + 1].cells[4], influence)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # T11: 2.1 AI Policy Statement (13x2)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    t = tables[11]
    policy = [
        "AI Policy for HealthWithSevgi ML Visualization Tool — Version 1.0, April 2026",

        "This policy establishes the governance framework for responsible development and "
        "educational deployment of the HealthWithSevgi ML Visualization Tool. It ensures that "
        "all AI-related activities adhere to principles of fairness, transparency, human "
        "oversight, privacy, safety, and accountability, in alignment with ISO/IEC 42001:2023 "
        "and the EU AI Act.",

        "This policy governs all AI/ML components as defined in Section 1.2, including: 8 "
        "supervised classification algorithms; the SHAP-based explainability engine; the "
        "subgroup fairness analysis system; all 20 built-in clinical datasets and user-uploaded "
        "CSV processing; the 7-step wizard pipeline; and the Docker-containerised deployment. "
        "Applies to all five team members across all five project sprints.",

        "We commit to equitable AI performance across all patient groups. Implementation: "
        "Step 7 subgroup fairness audit evaluates sensitivity across gender (Female/Male) and "
        "age bands (18–60, 61–75, 76+). 10pp gap triggers 'review' warning; 20pp or "
        "sensitivity <50% triggers 'action_needed' alert. Step 3 SMOTE addresses class "
        "imbalance with adaptive k_neighbors = max(1, min(5, min_class_count-1)). Three "
        "real-world AI failure case studies in Step 7 educate about consequences of unfair AI.",

        "We commit to making AI decisions interpretable for non-technical users. "
        "Implementation: Step 5 presents 7 metrics with colour-coded thresholds. Step 6 "
        "provides SHAP global feature importance and single-patient waterfall charts. 251 "
        "clinical name mappings translate raw variables to clinician-friendly terms. Clinical "
        "sense-check notes accompany top features.",

        "We commit to humans remaining the ultimate decision-makers. Implementation: Step 7 "
        "EU AI Act checklist item requires user confirmation of human oversight plan. Step 6 "
        "displays: 'A clinician must always decide whether and how to act on any AI "
        "prediction. This tool is an educational aid, not a diagnostic device.' No automated "
        "clinical actions are triggered.",

        "We commit to privacy-by-design. Implementation: All processing in-memory on FastAPI "
        "with OrderedDict LRU eviction (max 50 sessions/models). No database, no disk writes. "
        "Footer: 'Patient data is processed locally within this session. No patient data is "
        "stored or transmitted.' CORS restricted to localhost. PDF certificate contains only "
        "aggregate metrics — no raw patient data.",

        "We commit to preventing harmful outputs through technical safeguards. Implementation: "
        "Backend validates inputs (≥10 rows, ≤50 MB, ≤20 classes, ≥2 columns — HTTP 422 on "
        "failure). SHAP fallback chain: TreeExplainer → LinearExplainer → KernelExplainer → "
        "permutation importance. Step 5 low-sensitivity warning (<50%) and overfitting warning "
        "(train-test gap >10pp). Step 1: 'This is a learning tool. No real patient data "
        "leaves your browser session.'",

        "We commit to clear responsibility. Implementation: Governance roles mapped in "
        "Section 2.2. Sprint gates serve as management review (Clause 9.3). Git version "
        "control (101+ commits, 5 members) provides full audit trail. PDCA cycle: Ch.1–3 "
        "first submission (Plan) → implementation (Do) → feedback + peer audit (Check) → "
        "corrective actions (Act).",

        "ISO/IEC 42001:2023; EU AI Act Regulation 2024/1689 (educational alignment, Art. 6, "
        "Annex III); ISO/IEC 25059 (AI system quality); KVKK Law No. 6698 / EU GDPR (data "
        "protection alignment through privacy-by-design).",

        "Reviewed: (1) at project completion (May 2026); (2) when Step 7 detects sensitivity "
        "gap ≥20pp ('action_needed'); (3) when AIMS scope changes; (4) when regulations update "
        "(EU AI Act enforcement, KVKK amendments).",

        "Team: HealthWithSevgi, AI System Owner: Efe Çelik (Product Owner), Date: April 2026",
    ]
    for i, txt in enumerate(policy):
        set_cell_text(t.rows[i + 1].cells[1], txt)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # T12: 2.2 AI Governance Roles (6x4)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    t = tables[12]
    roles = [
        ("AI System Owner", "Product Owner", "Efe Çelik",
         "Overall accountability for AI system and ISO 42001 compliance. AI policy approval "
         "(Section 2.1); AIMS scope decisions (Section 1.2); final authority on risk "
         "acceptance; user story prioritisation; stakeholder communication."),
        ("Technical Lead", "Scrum Master /\nLead Developer", "Batuhan Bayazıt",
         "Technical oversight of all ML implementations and architecture. 8 algorithm "
         "implementations (MLService); SHAP engine quality across 4 explainer types; Docker "
         "deployment; CI/CD; session architecture design; code review; sprint facilitation."),
        ("Data Steward", "Developer\n(Backend)", "Berat Mert Gökkaya",
         "Data quality across 20 datasets (Section 3.1). Provenance documentation; dataset "
         "quality validation; bias/limitation documentation; Jira sprint management; backend "
         "ML pipeline implementation."),
        ("Ethics & UX\nReviewer", "UX Designer", "Burak Aydoğmuş",
         "UX design and ethical AI interface review. Clinical language clarity across all 7 "
         "Steps; Step 5 Results visualisation design; ensuring outputs support human oversight; "
         "UI accessibility; frontend implementation."),
        ("QA & Compliance\nLead", "QA / Documentation\nLead", "Berfin Duru Alkan",
         "Quality assurance and ISO 42001 compliance documentation. Test report generation; "
         "sprint progress reports; ISO report quality; peer audit coordination (Ch. 6); full "
         "pipeline test reports; cross-referencing report claims to code artefacts."),
    ]
    for i, (gov_role, scrum, name, resp) in enumerate(roles):
        set_cell_text(t.rows[i + 1].cells[0], gov_role)
        set_cell_text(t.rows[i + 1].cells[1], scrum)
        set_cell_text(t.rows[i + 1].cells[2], name)
        set_cell_text(t.rows[i + 1].cells[3], resp)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # T15: 3.1 Data Source Inventory (4x7 template — need to expand to 21 rows)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    t = tables[15]
    datasets = [
        ("1","Cardiology — Heart Failure","Heart Failure Clinical Records (Kaggle)","299","12","DEATH_EVENT (binary)","Small sample; ~32% death class imbalance; single-centre (Pakistan); limited demographics"),
        ("2","Cardiology — Stroke","Stroke Prediction Dataset (Kaggle)","5,110","10","stroke (binary)","Severe imbalance (<5% positive); self-reported data; socioeconomic proxies"),
        ("3","Cardiology — Arrhythmia","UCI Arrhythmia Dataset (UCI)","452","33","arrhythmia (binary)","High dimensionality; significant missing values; 1980s–90s data; binarised from 16 classes"),
        ("4","Radiology","NIH Chest X-Ray Metadata (Kaggle)","100,000+","4","Finding_Label (binary)","Metadata only (no images); NLP-extracted labels with noise; institutional bias"),
        ("5","Nephrology — CKD","UCI CKD Dataset (UCI)","400","15","classification (binary)","Significant missing values; single-centre (India); small sample for 15 features"),
        ("6","Oncology — Breast","Breast Cancer Wisconsin (UCI)","569","14","diagnosis (binary)","FNA-derived features; ~37% malignant; single institution; lacks demographics"),
        ("7","Oncology — Cervical","Cervical Cancer Risk Factors (UCI)","858","22","Biopsy (binary)","Extensive missing values; <10% positive; self-reported behavioural data; single country"),
        ("8","Neurology — Parkinson's","UCI Parkinson's Dataset (UCI)","195","17","status (binary)","Very small (n=195); ~75% PD positive; voice-only features; single session"),
        ("9","Endocrinology — Diabetes","Pima Indians Diabetes (Kaggle)","768","8","Outcome (binary)","Only Pima Indian women ≥21; zero values = missing; severely limits generalisability"),
        ("10","Endocrinology — Thyroid","UCI New Thyroid Dataset (UCI)","215","5","class (multiclass)","Very small (n=215); only 5 features; 1980s data; reference ranges may differ"),
        ("11","Hepatology — Liver","Indian Liver Patient (UCI)","583","10","Dataset (binary)","Single country (India); 75% male; ~72% positive; no aetiology info"),
        ("12","Mental Health — Depression","Depression Dataset (Kaggle)","~2,000","14","severity_class (binary)","Self-reported; oversimplifies severity; no DSM-5 criteria; weak provenance — used as pedagogical example"),
        ("13","Pulmonology — COPD","COPD EHR Dataset (PhysioNet)","~1,000","11","exacerbation (binary)","Credentialed access required; documentation bias; spirometry effort-dependent"),
        ("14","Haematology — Anaemia","Anaemia Detection Dataset (Kaggle)","~1,421","5","anemia_type (binary)","Only 5 features; no iron/ferritin/B12; binary — no subtype distinction"),
        ("15","Dermatology","HAM10000 Metadata (Harvard Dataverse)","10,015","3","dx_type (binary)","Metadata only; 3 features; ~10% melanoma; predominantly light-skinned (Vienna, Queensland)"),
        ("16","Ophthalmology — DR","Diabetic Retinopathy Debrecen (UCI)","1,151","19","severity_grade (binary)","Automated image features; single institution (Hungary); binarised from multi-grade"),
        ("17","Orthopaedics — Spine","Vertebral Column Dataset (UCI)","310","6","class (binary)","Only 6 features; small (n=310); no age/gender; examiner-dependent measurements"),
        ("18","ICU / Sepsis","PhysioNet Sepsis Challenge 2019","~40,000","16","SepsisLabel (binary)","~10% positive; missing values clinically meaningful; credentialed access required"),
        ("19","Obstetrics — Fetal","Cardiotocography Dataset (UCI)","2,126","12","fetal_health (multiclass)","~8% pathological; inter-observer variability; software-derived features"),
        ("20","Pharmacy — Readmission","Diabetes 130-US Hospitals (UCI)","100,000+","19","readmitted (multiclass)","1999–2008 data; US-specific; encounter-level; outdated treatment protocols"),
    ]

    # Fill existing rows first (row 0 = header, rows 1-3 = template data rows)
    for i in range(min(len(datasets), len(t.rows) - 1)):
        row = t.rows[i + 1]
        for j, val in enumerate(datasets[i]):
            set_cell_text(row.cells[j], val)

    # Add remaining rows
    template_data_row = t.rows[1]  # Use first data row as style reference
    for i in range(len(t.rows) - 1, len(datasets)):
        new_row = add_table_row(t, datasets[i])
        copy_row_style(template_data_row, new_row)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # T16: 3.2 Data Quality & Validation (5x2)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    t = tables[16]
    quality = [
        (
            "Quality Criteria",
            "• Minimum rows: ≥10 patients (MIN_ROWS=10)\n"
            "• Required columns: ≥2 (1 feature + 1 target)\n"
            "• File format: .csv only\n"
            "• Max file size: ≤50 MB (MAX_UPLOAD_MB=50)\n"
            "• Max target classes: ≤20 (MAX_TARGET_CLASSES=20)\n"
            "• Imbalance warning: ratio >1.5:1 triggers alert in Step 2\n"
            "• Target: selected via Column Mapper in Step 2\n"
            "• On failure: HTTP 422 with specific error → inline alert in Step 2 UI"
        ),
        (
            "Missing Value\nHandling",
            "• Median (default): Robust to outliers in clinical data (e.g., extreme "
            "creatinine in renal failure). Computed from training set only, applied to both "
            "sets to prevent leakage. Clinically preserves central tendency.\n"
            "• Mode: For categorical variables (gender, diagnosis codes). train_df.mode()."
            "iloc[0] with fallback to median.\n"
            "• Drop rows: Only when missingness is MCAR and dataset large enough. Risk: "
            "systematic missingness biases toward sicker patients who received more testing."
        ),
        (
            "Normalisation",
            "• Z-score (default): StandardScaler — zero mean, unit variance. Essential for "
            "KNN/SVM where magnitude affects distance. Intuitive for clinicians (standard "
            "deviations from mean).\n"
            "• Min-Max: MinMaxScaler — [0,1] range. For bounded features. Less robust to "
            "outliers.\n"
            "• None: Acceptable for tree-based models (DT, RF, XGBoost, LightGBM) which are "
            "scale-invariant. Risk: KNN/SVM dominated by high-magnitude features."
        ),
        (
            "Class Imbalance\n(SMOTE)",
            "• What: Generates synthetic minority samples by interpolating between k nearest "
            "minority neighbours.\n"
            "• Why needed: Clinical datasets often have severe imbalance (sepsis ~10%, stroke "
            "<5%). Without SMOTE, models optimise for majority class — high accuracy but 0% "
            "sensitivity for clinically important minority.\n"
            "• Adaptive k: max(1, min(5, min_class_count-1)) for small classes.\n"
            "• Training data only: Test set must reflect real-world distribution for honest "
            "metrics (ISO 42001 Annex A.6.4).\n"
            "• Step 3 displays: before/after class distribution comparison chart."
        ),
    ]
    for i, (topic, content) in enumerate(quality):
        set_cell_text(t.rows[i + 1].cells[0], topic, bold=True)
        set_cell_text(t.rows[i + 1].cells[1], content)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # T17: 3.3 Data Privacy & Provenance (4x2)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    t = tables[17]
    privacy = [
        (
            "Privacy\nArchitecture",
            "• In-memory processing: CSV → pandas DataFrame in server memory. No disk, no "
            "database, no external APIs.\n"
            "• Session-only lifecycle: UUID per session. MLService uses OrderedDict LRU "
            "eviction (max 50 sessions, 50 models). Auto-eviction on 51st entry.\n"
            "• Privacy notice (footer): 'Patient data is processed locally within this "
            "session. No patient data is stored or transmitted.'\n"
            "• CORS: localhost:5173 only (dev). Docker production: same-origin.\n"
            "• Certificate: aggregate metrics only — no raw patient data.\n"
            "• For real data under KVKK/GDPR: explicit consent (Art. 5/6), DPIA (Art. 35), "
            "encryption (AES-256/TLS 1.3), right-to-deletion, DPO, DPA, audit trail."
        ),
        (
            "Provenance\nTracking",
            "• Built-in datasets: Source URL, description, and license in specialty registry. "
            "Step 2 displays dataset attribution.\n"
            "• User uploads: filename, row/column counts, types, class distribution captured "
            "per session. No persistent lineage (educational scope).\n"
            "• Clinical deployment would need: Apache Atlas / OpenLineage, immutable provenance "
            "records, version control for dataset updates, source-to-prediction traceability."
        ),
        (
            "Bias in Data",
            "• Class imbalance: Stroke (<5%), Sepsis (~10%), Cervical Cancer (<10%), CTG (~8% "
            "pathological). Step 2 auto-warns at ratio >1.5:1. Step 3 offers SMOTE.\n"
            "• Demographic gaps: Pima diabetes = only Pima Indian women; ILPD = 75% male, "
            "India-only; HAM10000 = light-skinned (Austria/Australia).\n"
            "• Historical disparities: Missing labs may reflect under-testing of low-risk "
            "patients; Pharmacy dataset (1999–2008) reflects outdated protocols.\n"
            "• Step 7: Gender distribution chart vs. population norms (50/50); age-band "
            "analysis; subgroup performance table shows per-group sensitivity/specificity."
        ),
    ]
    for i, (topic, content) in enumerate(privacy):
        set_cell_text(t.rows[i + 1].cells[0], topic, bold=True)
        set_cell_text(t.rows[i + 1].cells[1], content)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # T18: 3.4 Third-Party Register (8x8 template)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    t = tables[18]
    deps = [
        ("scikit-learn","ML Library","≥1.4.0","BSD-3-Clause","No critical CVEs; NumFOCUS-sponsored; 60k+ stars","L","TensorFlow, PyTorch","Pin version; check PyPI releases each sprint"),
        ("XGBoost","ML Library","≥2.0.0","Apache 2.0","DMLC foundation; regular patches; enterprise adoption","L","LightGBM, CatBoost","Pin version; monitor GitHub advisories"),
        ("LightGBM","ML Library","≥4.3.0","MIT","Microsoft-maintained; requires libgomp1; no known CVEs","L","XGBoost, CatBoost","Pin version; monitor MSRC"),
        ("SHAP","Explainability","≥0.45.0","MIT","Academic origin (Lundberg); 18k+ citations; transitive deps","L","LIME, ELI5","Pin version; monitor API changes"),
        ("pandas","Data Processing","≥2.2.0","BSD-3-Clause","NumFOCUS; enterprise-grade; no critical vulns; 40k+ stars","L","Polars, Modin","Pin version; ecosystem stable"),
        ("imbalanced-learn","ML (SMOTE)","≥0.12.0","MIT","scikit-learn-contrib; well-tested SMOTE; quarterly releases","L","Custom SMOTE, ADASYN","Pin; follows sklearn cycle"),
        ("FastAPI","Web Framework","≥0.110.0","MIT","Active maintenance; Starlette/Pydantic; production-proven","L","Flask, Django REST","Pin version; monitor advisories"),
        ("ReportLab","PDF Generation","≥4.1.0","BSD","Enterprise PDF (20+ years); occasional CVE patches","L","WeasyPrint, FPDF2","Pin; scan with pip-audit"),
        ("React","Frontend","^18.3.0","MIT","Meta-maintained; dedicated security team; 220k+ stars","L","Vue, Svelte, Angular","Pin major; monitor security blog"),
        ("Vite","Build Tool","^5.3.0","MIT","Evan You / team; fast HMR; dev-only (not in runtime)","L","Webpack, Turbopack","Pin; monitor GitHub advisories"),
        ("Recharts","Charting","^2.12.0","MIT","React-specific on D3; ROC, PR, SHAP charts; 23k+ stars","L","@nivo, Plotly, Chart.js","Pin; monitor npm advisories"),
        ("Axios","HTTP Client","^1.7.0","MIT","Frontend API client; 120s timeout; /api baseURL; 100k+ stars","L","fetch API, ky","Pin; monitor npm advisories"),
        ("python:3.12-slim","Docker Base","3.12","PSF","Official image; regular security rebuilds; slim variant","L","python:3.12-alpine","Rebuild on advisories; scan with Trivy"),
        ("UCI ML Repository","Data Source","N/A","CC BY 4.0","Peer-reviewed; UC Irvine since 1987; benchmark standard","M","Kaggle, PhysioNet","Check for corrections/retractions"),
        ("Kaggle","Data Source","N/A","Various","Community-sourced; quality varies; cross-reference papers","M","UCI, govt data portals","Verify license per dataset"),
        ("PhysioNet","Data Source","N/A","PhysioNet License","Credentialed access; HIPAA de-identified; peer-reviewed","L","MIMIC alternatives","Monitor for corrections"),
        ("Harvard Dataverse","Data Source","N/A","CC BY 4.0","DOI-based; Harvard-maintained; persistent identifiers","L","ISIC Archive, Kaggle","Monitor via DOI resolver"),
    ]

    # Fill existing rows
    for i in range(min(len(deps), len(t.rows) - 1)):
        row = t.rows[i + 1]
        for j, val in enumerate(deps[i]):
            set_cell_text(row.cells[j], val)

    # Add extra rows
    template_data_row = t.rows[1]
    for i in range(len(t.rows) - 1, len(deps)):
        new_row = add_table_row(t, deps[i])
        copy_row_style(template_data_row, new_row)

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # Save
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
